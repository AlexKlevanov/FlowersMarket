import sys
import docx
import psycopg2
from PyQt6 import QtWidgets
from PyQt6.QtWidgets import *
from psycopg2 import Error
from PyQt6.QtCore import Qt

try:
    connection = psycopg2.connect(user="postgres",
                                  password="1234",
                                  host="127.0.0.1",
                                  port="5432",
                                  database="flowers market")

except (Exception, Error) as error:
    print("Ошибка при работе с PostgreSQL", error)
finally:
    if connection:
        #cursor.close()
        #connection.close()
        print("Соединение с PostgreSQL закрыто")




#Создание главного окна

class Window(QMainWindow):
    def __init__(self):
        super(Window, self).__init__()

        self.setWindowTitle("Цветочный магазин")
        self.setGeometry(300, 250, 350, 400)

        flowers_list = QtWidgets.QPushButton("Список цветов", self)
        flowers_list.setGeometry(120, 50, 120, 50)
        flowers_list.clicked.connect(self.Flowers)

        order = QtWidgets.QPushButton("Заказ", self)
        order.setGeometry(120, 110, 120, 50)
        order.clicked.connect(self.Order)

        warehouse = QtWidgets.QPushButton("Склад", self)
        warehouse.setGeometry(120, 170, 120, 50)
        warehouse.clicked.connect(self.Warehouse)

        shipment = QtWidgets.QPushButton("Поставка", self)
        shipment.setGeometry(120, 230, 120, 50)
        shipment.clicked.connect(self.Shipment)

        sales = QtWidgets.QPushButton("Продажи", self)
        sales.setGeometry(120, 290, 120, 50)
        sales.clicked.connect(self.Sales)

    #Создание окна со списком цветов.
    def Flowers(self):
        dlg1 = QDialog(self)
        dlg1.setWindowTitle("Список цветок")
        dlg1.setGeometry(120, 150, 700, 500)


        #Запрос в БД
        cursor = connection.cursor()
        insert_query = f"SELECT * FROM Цветок"
        cursor.execute(insert_query)
        connection.commit()
        Flist = cursor.fetchall()


        #Создание таблицы
        table = QTableWidget(dlg1)
        table.resize(800, 500)
        table.setColumnCount(4)
        table.setRowCount(len(Flist))
        table.setHorizontalHeaderLabels(['Код вида', 'Название', 'Количество на складе', 'Цена'])
        table.setColumnWidth(0, 100)
        table.setColumnWidth(1, 200)
        table.setColumnWidth(2, 200)
        table.setColumnWidth(3, 200)

        #Заполнение таблицы
        for i in range(len(Flist)):
            elem = Flist[i]
            for j in range(0, 4):
                table.setItem(i, j, QTableWidgetItem(str(elem[j])))

        dlg1.exec()

    #Создание окна для формирования заказа.
    def Order(self):
        dlg8 = QDialog(self)
        dlg8.setWindowTitle("Выбор")
        dlg8.resize(100, 100)

        layout = QVBoxLayout()
        dlg8.setLayout(layout)

        self.button_1 = QPushButton(self)
        self.button_1.setText("Оформление заказа")
        layout.addWidget(self.button_1, alignment=Qt.AlignmentFlag.AlignCenter)
        self.button_1.clicked.connect(self.Order_1)

        self.button_2 = QPushButton(self)
        self.button_2.setText("Добавление клиента")
        layout.addWidget(self.button_2, alignment=Qt.AlignmentFlag.AlignCenter)
        self.button_2.clicked.connect(self.Client)

        dlg8.exec()

    def Client(self):
        dlg9 = QDialog(self)
        dlg9.setWindowTitle("Клиент")
        dlg9.resize(200, 200)

        layout = QVBoxLayout()
        dlg9.setLayout(layout)

        self.number_client = QLabel(self)
        self.number_client.setText("Номер телефона клиента")
        self.Number_client = QLineEdit(self)
        layout.addWidget(self.number_client, alignment=Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(self.Number_client, alignment=Qt.AlignmentFlag.AlignCenter)

        self.name_client = QLabel(self)
        self.name_client.setText("ФИО клиента")
        self.Name_client = QLineEdit(self)
        layout.addWidget(self.name_client, alignment=Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(self.Name_client, alignment=Qt.AlignmentFlag.AlignCenter)

        self.adres_client = QLabel(self)
        self.adres_client.setText("Адрес клиента")
        self.Adres_client = QLineEdit(self)
        layout.addWidget(self.adres_client, alignment=Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(self.Adres_client, alignment=Qt.AlignmentFlag.AlignCenter)

        self.button_3 = QPushButton(self)
        self.button_3.setText("Добавить клиента")
        layout.addWidget(self.button_3, alignment=Qt.AlignmentFlag.AlignCenter)
        self.button_3.clicked.connect(self.client_add)

        dlg9.exec()
    def client_add(self):
        input_text_2 = self.Number_client.text()
        input_text_7 = self.Name_client.text()
        input_text_6 = self.Adres_client.text()

        cursor = connection.cursor()
        insert_query = f"INSERT INTO Клиент (Номер_телефона_клиента, ФИО_клиента, Адрес_клиента) VALUES ('{input_text_2}', '{input_text_7}', '{input_text_6}');"
        cursor.execute(insert_query)
        connection.commit()

        dlg11 = QMessageBox(self)
        dlg11.setWindowTitle("Добавлена новая запись")
        dlg11.setText("Клиент успешно добавлен в БД")
        button2 = dlg11.exec()
    def Order_1(self):
        dlg2 = QDialog(self)
        dlg2.setWindowTitle("Заказ")
        dlg2.resize(300, 300)

        self.number_employee = QLabel(self)
        self.number_employee.setText("Номер телефона сотрудника")
        self.Number_employee = QLineEdit(self)

        self.number_client = QLabel(self)
        self.number_client.setText("Номер телефона клиента")
        self.Number_client = QLineEdit(self)

        self.order_date = QLabel(self)
        self.order_date.setText("Дата заказа")
        self.Order_date = QCalendarWidget(self)

        self.flowers_type = QLabel(self)
        self.flowers_type.setText("Код вида цветка")
        self.Flowers_type = QComboBox(self)

        cursor = connection.cursor()
        insert_query = f"SELECT Код_вида FROM Цветок"
        cursor.execute(insert_query)
        connection.commit()
        code_flowers = cursor.fetchall()

        for i in range(len(code_flowers)):
            for j in code_flowers[i]:
                self.Flowers_type.addItem(str(j))

        self.flowers_count = QLabel(self)
        self.flowers_count.setText("Количество цветков")
        self.Flowers_count = QLineEdit(self)


        layout = QVBoxLayout()
        dlg2.setLayout(layout)


        layout.addWidget(self.number_employee, alignment=Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(self.Number_employee, alignment=Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(self.number_client, alignment=Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(self.Number_client, alignment=Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(self.order_date, alignment=Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(self.Order_date, alignment=Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(self.flowers_type, alignment=Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(self.Flowers_type, alignment=Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(self.flowers_count, alignment=Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(self.Flowers_count, alignment=Qt.AlignmentFlag.AlignCenter)


        self.order_add = QPushButton(self)
        self.order_add.setText("Сформировать заказ")
        layout.addWidget(self.order_add, alignment=Qt.AlignmentFlag.AlignCenter)
        self.order_add.clicked.connect(self.order_create)

        dlg2.exec()

    def order_create(self):
        input_text_1 = self.Number_employee.text()
        input_text_2 = self.Number_client.text()
        input_text_3 = self.Order_date.selectedDate()
        input_text_3 = input_text_3.toString(Qt.DateFormat.ISODate)
        input_text_4 = self.Flowers_type.currentText()
        input_text_5 = self.Flowers_count.text()

        cursor = connection.cursor()
        insert_query = f"SELECT Цена FROM Цветок WHERE Код_вида = '{int(input_text_4)}'"
        cursor.execute(insert_query)
        connection.commit()
        price = cursor.fetchall()
        price_1 = price[0]
        price_2 = price_1[0]

        cursor = connection.cursor()
        insert_query = f"INSERT INTO Заказ (Номер_телефона_сотрудника, Номер_телефона_клиента, Дата_заказа, Итоговая_стоимость) VALUES ('{input_text_1}', '{input_text_2}', '{input_text_3}', '{float(input_text_5) * price_2}');"
        cursor.execute(insert_query)
        connection.commit()

        cursor = connection.cursor()
        insert_query = f"SELECT Количество_цветков_на_складе FROM Цветок WHERE Код_вида = '{int(input_text_4)}'"
        cursor.execute(insert_query)
        connection.commit()
        count = cursor.fetchall()
        count_1 = count[0]
        count_2 = count_1[0]

        cursor = connection.cursor()
        insert_query = f"update Цветок set Количество_цветков_на_складе = '{int(count_2) - int(input_text_5)}' where Код_вида = '{input_text_4}'"
        cursor.execute(insert_query)
        connection.commit()

        cursor = connection.cursor()
        insert_query = f"select max(Номер_заказа) from Заказ;"
        cursor.execute(insert_query)
        connection.commit()
        max_count = cursor.fetchall()
        max_count_1 = max_count[0]
        max_count_2 = max_count_1[0]

        cursor = connection.cursor()
        insert_query = f"INSERT INTO Заказ_цветок (Номер_заказа, Количество_цветков_в_заказе, Код_вида) VALUES ('{max_count_2}', '{int(input_text_5)}', '{int(input_text_4)}')"
        cursor.execute(insert_query)
        connection.commit()

        a = str(float(input_text_5) * price_2)
        dlg10 = QMessageBox(self)
        dlg10.setWindowTitle("Добавлена новая запись")
        dlg10.setText("Заказ сформирован. Итоговая стоимость: " + a)
        button2 = dlg10.exec()

    #Создание окна для склада.
    def Warehouse(self):
        dlg3 = QDialog(self)
        dlg3.setWindowTitle("Склад")
        dlg3.resize(100, 200)

        layout1 = QVBoxLayout()
        dlg3.setLayout(layout1)

        self.warehouse_flowers_type = QLabel(self)
        self.warehouse_flowers_type.setText("Выберите вид цветка")
        self.Warehouse_flowers = QComboBox(self)
        layout1.addWidget(self.warehouse_flowers_type, alignment=Qt.AlignmentFlag.AlignCenter)
        layout1.addWidget(self.Warehouse_flowers, alignment=Qt.AlignmentFlag.AlignCenter)


        cursor = connection.cursor()
        insert_query = f"SELECT Название FROM Цветок"
        cursor.execute(insert_query)
        connection.commit()
        list_flowers = cursor.fetchall()

        for i in range(len(list_flowers)):
            for j in list_flowers[i]:
                self.Warehouse_flowers.addItem(j)

        self.warehouse_flowers_count = QLabel(self)
        self.warehouse_flowers_count.setText("Введите количество цветков, которое нужно вычесть")
        self.Warehouse_flowers_count = QLineEdit(self)
        layout1.addWidget(self.warehouse_flowers_count, alignment=Qt.AlignmentFlag.AlignCenter)
        layout1.addWidget(self.Warehouse_flowers_count, alignment=Qt.AlignmentFlag.AlignCenter)

        self.warehouse_actual = QPushButton("Ввести")
        layout1.addWidget(self.warehouse_actual, alignment=Qt.AlignmentFlag.AlignCenter)
        self.warehouse_actual.clicked.connect(self.Warehouse_actual)

        dlg3.exec()

    def Warehouse_actual(self):
        input_text_1 = self.Warehouse_flowers.currentText()
        input_text_2 = self.Warehouse_flowers_count.text()

        cursor = connection.cursor()
        insert_query = f"SELECT Количество_цветков_на_складе FROM Цветок WHERE Название = '{input_text_1}'"
        cursor.execute(insert_query)
        connection.commit()
        count = cursor.fetchall()

        count_1 = count[0]
        count_2 = count_1[0]

        cursor = connection.cursor()
        insert_query = f"UPDATE Цветок SET Количество_цветков_на_складе = '{int(count_2) - int(input_text_2)}' WHERE Название = '{input_text_1}'"
        cursor.execute(insert_query)
        connection.commit()

        dlg6 = QMessageBox(self)
        dlg6.setWindowTitle("Добавлена новая запись")
        dlg6.setText("Количество цветков на складе успешно измненено")
        button2 = dlg6.exec()


    #Создание окна для поставок.
    def Shipment(self):
        dlg4 = QDialog(self)
        dlg4.setWindowTitle("Поставка")
        dlg4.resize(300, 300)

        layout = QVBoxLayout()
        dlg4.setLayout(layout)

        self.shipment_code = QLabel(self)
        self.shipment_code.setText("Введите код поставки")
        layout.addWidget(self.shipment_code, alignment=Qt.AlignmentFlag.AlignCenter)

        self.Shipment_code = QLineEdit(self)
        layout.addWidget(self.Shipment_code, alignment=Qt.AlignmentFlag.AlignCenter)

        self.shipment_date = QLabel(self)
        self.shipment_date.setText("Укажите дату поставки")
        layout.addWidget(self.shipment_date, alignment=Qt.AlignmentFlag.AlignCenter)

        self.Shipment_date = QCalendarWidget(self)
        layout.addWidget(self.Shipment_date, alignment=Qt.AlignmentFlag.AlignCenter)

        self.shipment_flowers_type = QLabel(self)
        self.shipment_flowers_type.setText("Выберите код вида")
        layout.addWidget(self.shipment_flowers_type, alignment=Qt.AlignmentFlag.AlignCenter)

        self.Shipment_flowers_type = QComboBox(self)
        layout.addWidget(self.Shipment_flowers_type, alignment=Qt.AlignmentFlag.AlignCenter)

        cursor = connection.cursor()
        insert_query = f"SELECT Код_вида FROM Цветок"
        cursor.execute(insert_query)
        connection.commit()
        list_flowers = cursor.fetchall()

        for i in range(len(list_flowers)):
            for j in list_flowers[i]:
                self.Shipment_flowers_type.addItem(str(j))

        self.shipment_flowers_count = QLabel(self)
        self.shipment_flowers_count.setText("Введите количество цветков")
        layout.addWidget(self.shipment_flowers_count, alignment=Qt.AlignmentFlag.AlignCenter)

        self.Shipment_flowers_count = QLineEdit(self)
        layout.addWidget(self.Shipment_flowers_count, alignment=Qt.AlignmentFlag.AlignCenter)

        button = QPushButton("Добавить поставку")
        button.clicked.connect(self.Shipment_add)
        layout.addWidget(button)
        dlg4.exec()

    def Shipment_add(self):
        input_text_1 = self.Shipment_code.text()
        input_text_2 = self.Shipment_date.selectedDate()
        input_text_2 = input_text_2.toString(Qt.DateFormat.ISODate)
        input_text_3 = self.Shipment_flowers_count.text()
        input_text_4 = self.Shipment_flowers_type.currentText()

        cursor = connection.cursor()
        insert_query = f"INSERT INTO Поставка (Дата_поставки, Код_поставки) VALUES ('{input_text_2}', '{int(input_text_1)}')"
        cursor.execute(insert_query)
        connection.commit()

        cursor = connection.cursor()
        insert_query = f"INSERT INTO Состав_поставки (Код_поставки, Количество_цветков_в_поставке, Код_вида) VALUES ('{int(input_text_1)}', '{int(input_text_3)}', '{int(input_text_4)}')"
        cursor.execute(insert_query)
        connection.commit()

        cursor = connection.cursor()
        insert_query = f"SELECT Количество_цветков_на_складе FROM Цветок WHERE Код_вида = '{input_text_4}'"
        cursor.execute(insert_query)
        connection.commit()
        count = cursor.fetchall()
        count_1 = count[0]
        count_2 = count_1[0]
        print(count_2)
        cursor = connection.cursor()
        insert_query = f"UPDATE Цветок SET Количество_цветков_на_складе = '{int(count_2) + int(input_text_3)}' WHERE Код_вида = '{int(input_text_4)}'"
        cursor.execute(insert_query)
        connection.commit()

        dlg7 = QMessageBox(self)
        dlg7.setWindowTitle("Добавлена новая запись")
        dlg7.setText("Запись успешно добавлена в БД")
        button3 = dlg7.exec()

    #Создание окна для продаж.
    def Sales(self):
        dlg5 = QDialog(self)
        dlg5.setWindowTitle("Продажи")

        layout = QVBoxLayout()
        dlg5.setLayout(layout)

        self.sales_date = QLabel(self)
        self.sales_date.setText("Выберите начальную дату периода")
        self.Sale_date = QCalendarWidget(self)
        layout.addWidget(self.sales_date, alignment=Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(self.Sale_date, alignment=Qt.AlignmentFlag.AlignCenter)

        self.sales_date_2 = QLabel(self)
        self.sales_date_2.setText("Выберите конечную дату периода")
        self.Sale_date_2 = QCalendarWidget(self)
        layout.addWidget(self.sales_date_2, alignment=Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(self.Sale_date_2, alignment=Qt.AlignmentFlag.AlignCenter)

        self.button_sales = QPushButton(self)
        self.button_sales.setText("Сформировать отчёт")
        layout.addWidget(self.button_sales, alignment=Qt.AlignmentFlag.AlignCenter)
        self.button_sales.clicked.connect(self.get_rep)

        dlg5.exec()

    def get_rep(self):
        dlg6 = QDialog(self)
        dlg6.setWindowTitle("Отчёт по продажам")
        dlg6.resize(370, 300)


        input_text_1 = self.Sale_date.selectedDate()
        input_text_1 = input_text_1.toString(Qt.DateFormat.ISODate)
        input_text_2 = self.Sale_date_2.selectedDate()
        input_text_2 = input_text_2.toString(Qt.DateFormat.ISODate)

        cursor = connection.cursor()
        insert_query = f"SELECT Заказ_цветок.Код_вида, SUM(Заказ_цветок.Количество_цветков_в_заказе) AS Количество_проданных FROM Заказ_цветок JOIN Заказ ON Заказ_цветок.Номер_заказа = Заказ.Номер_заказа WHERE Заказ.Дата_заказа BETWEEN '{input_text_1}' AND '{input_text_2}' GROUP BY Заказ_цветок.Код_вида"
        cursor.execute(insert_query)
        connection.commit()
        report = cursor.fetchall()

        table = QTableWidget(dlg6)
        table.resize(370, 300)
        table.setColumnCount(2)
        table.setRowCount(len(report))
        table.setHorizontalHeaderLabels(['Код_вида', 'Количество_проданных'])
        table.setColumnWidth(0, 100)
        table.setColumnWidth(1, 250)

        #Заполнение таблицы
        for i in range(len(report)):
            elem = report[i]
            for j in range(0, 2):
                table.setItem(i, j, QTableWidgetItem(str(elem[j])))


        puth = f"C:\\Users\\kleva\\OneDrive\\Рабочий стол\\Отчёт_по_продажам_за_период_{input_text_1}-{input_text_2}.docx"
        doc = docx.Document()
        doc.add_heading(f" Отчёт по продажам за период {input_text_1} - {input_text_2} \n")
        table_doc = doc.add_table(rows=len(report) + 1, cols=2)
        table_doc.style = 'Table Grid'
        row = table_doc.rows[0]
        row.cells[0].text = 'Код вида'
        row.cells[1].text = 'Количество проданных'
        for i in range(1, len(report) + 1):
            element = report[i - 1]
            row_1 = table_doc.rows[i]
            for j in range(0, 2):
                row_1.cells[j].text = str(element[j])



        doc.save(puth)
        dlg6.exec()


app = QApplication(sys.argv)
window = Window()
window.show()
app.exec()

